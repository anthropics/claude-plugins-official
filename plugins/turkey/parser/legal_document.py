"""Provider-independent parsing of legal documents stored on disk."""
from __future__ import annotations

from dataclasses import dataclass, field
from enum import Enum
from html.parser import HTMLParser
from pathlib import Path
from typing import Protocol, runtime_checkable


class LegalDocumentType(str, Enum):
    PDF = "pdf"
    DOCX = "docx"
    HTML = "html"
    TXT = "txt"
    COURT_DECISION = "court-decision"
    LAW = "law"
    REGULATION = "regulation"
    COMMUNIQUE = "communique"
    PRECEDENT = "precedent"


@dataclass(frozen=True)
class ParsedLegalDocument:
    path: Path
    title: str
    text: str
    document_type: LegalDocumentType
    metadata: dict[str, str] = field(default_factory=dict)


@runtime_checkable
class LegalDocumentParser(Protocol):
    """Parses one local file without fetching from a provider or network."""

    def supports(self, path: Path) -> bool: ...

    def parse(
        self, path: Path, document_type: LegalDocumentType | None = None
    ) -> ParsedLegalDocument: ...


class _HTMLTextExtractor(HTMLParser):
    def __init__(self) -> None:
        super().__init__()
        self.parts: list[str] = []

    def handle_data(self, data: str) -> None:
        self.parts.append(data)


def infer_document_type(path: Path) -> LegalDocumentType:
    name = path.stem.lower()
    for marker, document_type in (
        ("mahkeme", LegalDocumentType.COURT_DECISION),
        ("karar", LegalDocumentType.COURT_DECISION),
        ("ictihat", LegalDocumentType.PRECEDENT),
        ("kanun", LegalDocumentType.LAW),
        ("yonetmelik", LegalDocumentType.REGULATION),
        ("teblig", LegalDocumentType.COMMUNIQUE),
    ):
        if marker in name:
            return document_type
    extension_types = {
        ".pdf": LegalDocumentType.PDF,
        ".docx": LegalDocumentType.DOCX,
        ".html": LegalDocumentType.HTML,
        ".htm": LegalDocumentType.HTML,
        ".txt": LegalDocumentType.TXT,
    }
    try:
        return extension_types[path.suffix.lower()]
    except KeyError as error:
        raise ValueError(f"Unsupported legal document format: {path.suffix or path.name}") from error


class TextLegalDocumentParser:
    def supports(self, path: Path) -> bool:
        return path.suffix.lower() == ".txt"

    def parse(
        self, path: Path, document_type: LegalDocumentType | None = None
    ) -> ParsedLegalDocument:
        return ParsedLegalDocument(
            path=path,
            title=path.stem,
            text=path.read_text(encoding="utf-8"),
            document_type=document_type or infer_document_type(path),
        )


class HtmlLegalDocumentParser:
    def supports(self, path: Path) -> bool:
        return path.suffix.lower() in {".html", ".htm"}

    def parse(
        self, path: Path, document_type: LegalDocumentType | None = None
    ) -> ParsedLegalDocument:
        extractor = _HTMLTextExtractor()
        extractor.feed(path.read_text(encoding="utf-8"))
        return ParsedLegalDocument(
            path=path,
            title=path.stem,
            text=" ".join(part.strip() for part in extractor.parts if part.strip()),
            document_type=document_type or infer_document_type(path),
        )


class PdfLegalDocumentParser:
    def supports(self, path: Path) -> bool:
        return path.suffix.lower() == ".pdf"

    def parse(
        self, path: Path, document_type: LegalDocumentType | None = None
    ) -> ParsedLegalDocument:
        try:
            from pypdf import PdfReader
        except ImportError as error:
            raise RuntimeError("PDF parsing requires the optional 'pypdf' dependency") from error
        text = "\n".join(page.extract_text() or "" for page in PdfReader(path).pages)
        return ParsedLegalDocument(path, path.stem, text, document_type or infer_document_type(path))


class DocxLegalDocumentParser:
    def supports(self, path: Path) -> bool:
        return path.suffix.lower() == ".docx"

    def parse(
        self, path: Path, document_type: LegalDocumentType | None = None
    ) -> ParsedLegalDocument:
        try:
            from docx import Document
        except ImportError as error:
            raise RuntimeError("DOCX parsing requires the optional 'python-docx' dependency") from error
        text = "\n".join(paragraph.text for paragraph in Document(path).paragraphs)
        return ParsedLegalDocument(path, path.stem, text, document_type or infer_document_type(path))


class LegalDocumentParserRegistry:
    """Selects a local parser by file extension; it has no provider dependency."""

    def __init__(self, parsers: list[LegalDocumentParser] | None = None):
        self._parsers = parsers or [
            TextLegalDocumentParser(),
            HtmlLegalDocumentParser(),
            PdfLegalDocumentParser(),
            DocxLegalDocumentParser(),
        ]

    def parse(
        self, path: Path, document_type: LegalDocumentType | None = None
    ) -> ParsedLegalDocument:
        for parser in self._parsers:
            if parser.supports(path):
                return parser.parse(path, document_type)
        raise ValueError(f"Unsupported legal document format: {path.suffix or path.name}")
